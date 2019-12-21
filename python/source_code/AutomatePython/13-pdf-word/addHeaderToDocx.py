import docx

doc = docx.Document()
doc.add_heading('Head 0', 0)
doc.add_heading('Head 1', 1)
doc.add_heading('Head 2', 2)
doc.add_heading('Head 3', 3)
doc.add_heading('Head 4', 4)

doc.save('headings.docx')
