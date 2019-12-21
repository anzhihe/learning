import docx

# create a docment
doc = docx.Document('guestTemplete.docx')
paragraphNum = len(doc.paragraphs)
guestDoc = docx.Document()

# open guests.txt file 
with open('guests.txt') as f:
    for guest in f.readlines():
        paraSeq = 0  
        # add each guest to the document    
        for p in doc.paragraphs:
            paraSeq += 1  
            if paraSeq == 2:
                p.add_run(guest)      
            if paraSeq == paragraphNum:
                p.runs[len(p.runs)-1].add_break(docx.text.run.WD_BREAK.PAGE)
            # set text and style
            guestDoc.add_paragraph(p.text, p.style)

# save file
guestDoc.save('guestsInvitaion.docx')