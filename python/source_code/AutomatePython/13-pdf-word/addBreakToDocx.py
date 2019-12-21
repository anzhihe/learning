import docx

doc = docx.Document()
doc.add_paragraph('This is on the first page!')
doc.paragraphs[0].runs[0].add_break(docx.text.run.WD_BREAK.PAGE)
doc.add_paragraph('This is on the second page!')

doc.save('twoPage.docx')