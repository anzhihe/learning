import docx

doc = docx.Document('demo.docx')
print(len(doc.paragraphs))
'''
7
'''

for p in doc.paragraphs:
	print(str(len(p.text)).ljust(8), p.text)

'''	
14       Document Title
48       A plain paragraph with some bold and some italic
16       Heading, level 1
13       Intense quote
28       first item in unordered list
26       first item in ordered list
1        
'''

for r in doc.paragraphs[1].runs:
	print(str(len(r.text)).ljust(8), r.text)

'''	
22       A plain paragraph with
6         some 
4        bold
10        and some 
6        italic
'''