from docx import Document
# $ pip3 install python-docx
from pathlib import Path, PurePath

# word文件所在路径
word_files_path = '/Users/edz/Desktop/效率专栏/文章2/word样例文件'

# 取得该目录下所有的docx格式文件
p = Path(word_files_path)
files = [x for x in p.iterdir() if PurePath(x).match('*.docx')]
# 实例化文档对象
doc = Document()

def merge_files(docx_files: list):
    '''
    保留文件原有格式进行合并
    '''
    # 遍历每个文件
    for file in sorted(docx_files):
        another_doc = Document(file)
        # 遍历每一页
        for word_body in another_doc.element.body:
            # 合并内容到新的word文档
            doc.element.body.append(word_body)

    doc.save(Path(word_files_path, 'new.docx'))

def merge_without_format(docx_files: list):
    '''
    只获取内容进行合并
    '''
    # 遍历每个文件
    for docx_file in sorted(docx_files):
        another_doc = Document(docx_file)
        # 获取每个文件的所有“段落”
        paras = another_doc.paragraphs
        # 获取所有段落的文字内容
        # paras_content = [para.text for para in paras]
        for para in paras:
            # 为新的word文件创建一个新段落
            newpar = doc.add_paragraph('')
            # 将提取的内容写入新的文本段落中
            newpar.add_run(para.text)

    # 所有文件合并完成后在指定路径进行保存
    doc.save(Path(word_files_path, 'new.docx'))
        

# 调用函数
merge_without_format(files)

